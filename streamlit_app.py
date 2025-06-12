import streamlit as st
import os
import re
import hashlib
from typing import List, Dict, Tuple
import pdfplumber
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.vectorstores import InMemoryVectorStore
from langchain_groq import ChatGroq
from langchain_core.prompts import ChatPromptTemplate
from langchain.embeddings import HuggingFaceEmbeddings
import pandas as pd
from langchain_core.documents import Document

from google.cloud import vision
from google.oauth2 import service_account
import json
import fitz
from PIL import Image
import io

from datetime import datetime
import xlsxwriter
from io import BytesIO

from docx import Document as DocxDocument
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

st.set_page_config(
    page_title="Financial Underwriting Assistant",
    page_icon="💰",
    layout="wide"
)

st.title("💰 Financial Underwriting Assitant")

def get_income_multiplier(age: int, policy_type: str) -> int:
    if policy_type.lower() == "term":
        if 18 <= age <= 30:
            return 25
        elif 31 <= age <= 35:
            return 25
        elif 36 <= age <= 40:
            return 20
        elif 41 <= age <= 45:
            return 15
        elif 46 <= age <= 50:
            return 12
        elif 51 <= age <= 55:
            return 10
        elif age >= 56:
            return 5
    else:
        if 18 <= age <= 30:
            return 35
        elif 31 <= age <= 35:
            return 30
        elif 36 <= age <= 40:
            return 25
        elif 41 <= age <= 45:
            return 20
        elif 46 <= age <= 50:
            return 15
        elif 51 <= age <= 65:
            return 10
        elif age > 65:
            return 6
    
    return 10

class PIIShield:
    def __init__(self):
        self.pii_patterns = {
            'pan_card': r'\b[A-Z]{5}[-\s]?[0-9]{4}[-\s]?[A-Z]\b',
            'tan': r'\b[A-Z]{4}[-\s]?[0-9]{5}[-\s]?[A-Z]\b',
            'aadhaar': r'\b\d{4}\s?\d{4}\s?\d{4}\b',
            'account_number': r'\b\d{9,18}\b',
            'phone': r'\b(?:\+91[-.\s]?)?[6-9]\d{9}\b',
            'email': r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
            'address': r'\b(?:house|flat|plot|door)\s*(?:no\.?|number)?\s*[0-9A-Za-z\-\/]+\b',
            'ifsc': r'\b[A-Z]{4}0[A-Z0-9]{6}\b'
        }
        
        self.replacement_map = {}
        self.anonymization_enabled = True
    
    def anonymize_text(self, text: str) -> str:
        if not self.anonymization_enabled:
            return text
            
        anonymized_text = text
        
        for pii_type, pattern in self.pii_patterns.items():
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                original = match.group()
                
                if original not in self.replacement_map:
                    self.replacement_map[original] = "###########"
                
                anonymized_text = anonymized_text.replace(original, self.replacement_map[original])
        
        return anonymized_text
    
    def get_pii_summary(self) -> Dict[str, int]:
        pii_summary = {}
        for original, anonymized in self.replacement_map.items():
            for pii_type, pattern in self.pii_patterns.items():
                if re.match(pattern, original, re.IGNORECASE):
                    pii_summary[pii_type] = pii_summary.get(pii_type, 0) + 1
                    break
        return pii_summary

class FinancialDataExtractor:
    def __init__(self):
        self.extracted_data = {
            'personal_info': [],
            'income_details': [],
            'investment_details': [],
            'bank_details': [],
            'loan_details': [],
            'tax_details': [],
            'tables_data': [],
            'raw_text_data': [],
            'salary_credits': []
        }
        
        self.financial_patterns = {
            'gross_salary': r'(?:gross\s*gross|ctc)[\s:]*[₹$]?\s*([0-9,]+\.?[0-9]*)',
            'investment_amount': r'(?:investment|invested|amount)[\s:]*[₹$]?\s*([0-9,]+\.?[0-9]*)',
            'balance': r'(?:balance|closing\s*balance)[\s:]*[₹$]?\s*([0-9,]+\.?[0-9]*)',
            'credit_limit': r'(?:credit\s*limit)[\s:]*[₹$]?\s*([0-9,]+\.?[0-9]*)',
            'emi': r'(?:emi|monthly\s*installment)[\s:]*[₹$]?\s*([0-9,]+\.?[0-9]*)',
            'tax_paid': r'(?:tax\s*paid|income\s*tax)[\s:]*[₹$]?\s*([0-9,]+\.?[0-9]*)'
        }
    
    def extract_from_documents(self, documents):
        for doc in documents:
            content = doc.page_content
            metadata = doc.metadata
            
            if metadata.get("type") == "table":
                self.extract_table_data(content, metadata)
            else:
                self.extract_text_data(content, metadata)
        
        return self.extracted_data
    
    def extract_table_data(self, content, metadata):
        lines = content.split('\n')
        table_data = []
        
        for line in lines:
            if line.strip() and not line.startswith('---'):
                amounts = re.findall(r'[₹$]?\s*([0-9,]+\.?[0-9]*)', line)
                if amounts:
                    table_data.append({
                        'source': metadata.get('source', ''),
                        'page': metadata.get('page', ''),
                        'table_number': metadata.get('table_number', ''),
                        'content': line.strip(),
                        'extracted_amounts': amounts
                    })
        
        if table_data:
            self.extracted_data['tables_data'].extend(table_data)
    
    def extract_text_data(self, content, metadata):
        content_lower = content.lower()
        
        for pattern_name, pattern in self.financial_patterns.items():
            matches = re.finditer(pattern, content_lower, re.IGNORECASE)
            for match in matches:
                extracted_item = {
                    'source': metadata.get('source', ''),
                    'page': metadata.get('page', ''),
                    'type': pattern_name,
                    'value': match.group(1),
                    'context': content[max(0, match.start()-50):match.end()+50].strip()
                }
                
                if any(keyword in content_lower for keyword in ['salary', 'income', 'pay']):
                    self.extracted_data['income_details'].append(extracted_item)
                elif any(keyword in content_lower for keyword in ['investment', 'mutual fund', 'sip']):
                    self.extracted_data['investment_details'].append(extracted_item)
                elif any(keyword in content_lower for keyword in ['bank', 'account', 'balance']):
                    self.extracted_data['bank_details'].append(extracted_item)
                elif any(keyword in content_lower for keyword in ['loan', 'emi', 'credit']):
                    self.extracted_data['loan_details'].append(extracted_item)
                elif any(keyword in content_lower for keyword in ['tax', 'itr']):
                    self.extracted_data['tax_details'].append(extracted_item)
        
        self.extracted_data['raw_text_data'].append({
            'source': metadata.get('source', ''),
            'page': metadata.get('page', ''),
            'content': content[:500] + '...' if len(content) > 500 else content
        })

def create_excel_export(extracted_data, filename="financial_data_export.xlsx"):
    buffer = BytesIO()
    
    with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
        workbook = writer.book
        
        header_format = workbook.add_format({
            'bold': True,
            'bg_color': '#4472C4',
            'font_color': 'white',
            'border': 1
        })
        
        cell_format = workbook.add_format({
            'border': 1,
            'text_wrap': True
        })
        
        amount_format = workbook.add_format({
            'border': 1,
            'num_format': '#,##0.00'
        })
        
        summary_data = []
        for category, items in extracted_data.items():
            if category != 'raw_text_data' and items:
                summary_data.append({
                    'Category': category.replace('_', ' ').title(),
                    'Items Count': len(items),
                    'Description': f"Contains {len(items)} extracted items"
                })
        
        if summary_data:
            summary_df = pd.DataFrame(summary_data)
            summary_df.to_excel(writer, sheet_name='Summary', index=False)
            worksheet = writer.sheets['Summary']
            worksheet.set_column('A:C', 20)
        
        for category, items in extracted_data.items():
            if not items or category == 'raw_text_data':
                continue
                
            sheet_name = category.replace('_', ' ').title()[:31]
            
            if category == 'tables_data':
                df_data = []
                for item in items:
                    df_data.append({
                        'Source': item.get('source', ''),
                        'Page': item.get('page', ''),
                        'Table Number': item.get('table_number', ''),
                        'Content': item.get('content', ''),
                        'Extracted Amounts': ', '.join(item.get('extracted_amounts', []))
                    })
                df = pd.DataFrame(df_data)
            else:
                df_data = []
                for item in items:
                    df_data.append({
                        'Source': item.get('source', ''),
                        'Page': item.get('page', ''),
                        'Type': item.get('type', ''),
                        'Value': item.get('value', ''),
                        'Context': item.get('context', '')
                    })
                df = pd.DataFrame(df_data)
            
            if not df.empty:
                df.to_excel(writer, sheet_name=sheet_name, index=False)
                worksheet = writer.sheets[sheet_name]
                
                for col_num, col_name in enumerate(df.columns):
                    if col_name in ['Value', 'Extracted Amounts']:
                        worksheet.set_column(col_num, col_num, 15, amount_format)
                    elif col_name in ['Context', 'Content']:
                        worksheet.set_column(col_num, col_num, 40, cell_format)
                    else:
                        worksheet.set_column(col_num, col_num, 20, cell_format)
                
                for col_num, _ in enumerate(df.columns):
                    worksheet.write(0, col_num, df.columns[col_num], header_format)
        
        if extracted_data.get('raw_text_data'):
            raw_data = extracted_data['raw_text_data'][:100]
            df_raw = pd.DataFrame(raw_data)
            df_raw.to_excel(writer, sheet_name='Raw Text Data', index=False)
            worksheet = writer.sheets['Raw Text Data']
            worksheet.set_column('A:C', 30, cell_format)
    
    buffer.seek(0)
    return buffer

pii_shield = PIIShield()
financial_extractor = FinancialDataExtractor()

def setup_vision_client():
    try:
        api_key = st.session_state.get('google_vision_api_key')
        
        if not api_key:
            st.error("⚠️ Please enter your Google Vision API key in the sidebar.")
            return None
            
        client = vision.ImageAnnotatorClient(
            client_options={"api_key": api_key}
        )
        return client
    except Exception as e:
        st.error(f"Error setting up Google Vision client: {str(e)}")
        return None

def pdf_to_images(pdf_path):
    try:
        doc = fitz.open(pdf_path)
        images = []
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            img_data = pix.tobytes("png")
            images.append({
                "data": img_data,
                "page": page_num + 1
            })
        
        doc.close()
        return images
    except Exception as e:
        st.error(f"Error converting PDF to images: {str(e)}")
        return []

def extract_text_with_vision(pdf_path):
    vision_client = setup_vision_client()
    if not vision_client:
        return []
    
    try:
        images = pdf_to_images(pdf_path)
        documents = []
        
        for img_info in images:
            image = vision.Image(content=img_info["data"])
            
            response = vision_client.text_detection(image=image)
            
            if response.error.message:
                st.error(f"Vision API error: {response.error.message}")
                continue
            
            texts = response.text_annotations
            if texts:
                extracted_text = texts[0].description
                
                doc = Document(
                    page_content=extracted_text,
                    metadata={
                        "source": pdf_path,
                        "page": img_info["page"],
                        "type": "scanned_text",
                        "extraction_method": "google_vision"
                    }
                )
                documents.append(doc)
        
        return documents
        
    except Exception as e:
        st.error(f"Error with Google Vision API: {str(e)}")
        return []

def extract_tables_with_vision(pdf_path):
    vision_client = setup_vision_client()
    if not vision_client:
        return []
    
    try:
        images = pdf_to_images(pdf_path)
        table_documents = []
        
        for img_info in images:
            image = vision.Image(content=img_info["data"])
            
            response = vision_client.document_text_detection(image=image)
            
            if response.error.message:
                st.error(f"Vision API error: {response.error.message}")
                continue
            
            if response.full_text_annotation:
                tables = extract_table_structures(response.full_text_annotation, img_info["page"])
                table_documents.extend(tables)
        
        return table_documents
        
    except Exception as e:
        st.error(f"Error extracting tables with Google Vision API: {str(e)}")
        return []

def extract_table_structures(full_text_annotation, page_num):
    tables = []
    
    blocks = full_text_annotation.pages[0].blocks if full_text_annotation.pages else []
    
    for block_idx, block in enumerate(blocks):
        if is_table_block(block):
            table_text = extract_block_text(block)
            table_csv = convert_to_csv_format(table_text)
            
            doc = Document(
                page_content=table_csv,
                metadata={
                    "page": page_num,
                    "type": "table_data",
                    "extraction_method": "google_vision_table",
                    "block_index": block_idx,
                    "format": "csv"
                }
            )
            tables.append(doc)
    
    return tables

def is_table_block(block):
    text = extract_block_text(block)
    lines = text.strip().split('\n')
    
    if len(lines) < 2:
        return False
    
    column_counts = []
    for line in lines:
        columns = len([col for col in line.split() if col.strip()])
        if columns > 1:
            column_counts.append(columns)
    
    if len(column_counts) >= 2:
        avg_cols = sum(column_counts) / len(column_counts)
        consistent_cols = sum(1 for count in column_counts if abs(count - avg_cols) <= 1)
        return consistent_cols / len(column_counts) >= 0.7
    
    return False

def extract_block_text(block):
    text_lines = []
    for paragraph in block.paragraphs:
        line_text = ""
        for word in paragraph.words:
            word_text = "".join([symbol.text for symbol in word.symbols])
            line_text += word_text + " "
        text_lines.append(line_text.strip())
    return "\n".join(text_lines)

def convert_to_csv_format(table_text):
    """Convert table text to CSV format"""
    lines = table_text.strip().split('\n')
    csv_lines = []
    
    for line in lines:
        columns = [col.strip() for col in line.split('  ') if col.strip()]
        if columns:
            escaped_columns = []
            for col in columns:
                if ',' in col or '"' in col:
                    escaped_col = col.replace('"', '""')
                    col = f'"{escaped_col}"'
                escaped_columns.append(col)
            csv_lines.append(','.join(escaped_columns))
    
    return '\n'.join(csv_lines)

def extract_content_with_vision(pdf_path, extract_tables=True):
    documents = []
    
    text_docs = extract_text_with_vision(pdf_path)
    documents.extend(text_docs)
    
    if extract_tables:
        table_docs = extract_tables_with_vision(pdf_path)
        documents.extend(table_docs)
        
        if table_docs:
            st.success(f"✅ Extracted {len(table_docs)} tables from the document")
    
    return documents

def is_scanned_pdf(pdf_path):
    try:
        with pdfplumber.open(pdf_path) as pdf:
            total_text = ""
            total_chars = 0
            pages_checked = min(3, len(pdf.pages))
            
            for page in pdf.pages[:pages_checked]:
                text = page.extract_text() or ""
                total_text += text
                total_chars += len(text.strip())
            
            
            avg_chars_per_page = total_chars / pages_checked if pages_checked > 0 else 0
            return avg_chars_per_page < 50
    except:
        return True

def extract_tables_from_pdf(file_path):
    document_content = []
    
    with pdfplumber.open(file_path) as pdf:
        for page_num, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            if text.strip():
                document_content.append({
                    "content": text,
                    "page": page_num + 1,
                    "type": "text"
                })
            
            tables = page.extract_tables()
            for table_num, table in enumerate(tables):
                if table:
                    df = pd.DataFrame(table)
                    
                    if not df.empty:
                        headers = []
                        if len(df.columns) > 0:
                            if not pd.isna(df.iloc[0]).all() and not all(x is None for x in df.iloc[0]):
                                headers = [str(h).strip() if h is not None else f"Column_{i}" 
                                          for i, h in enumerate(df.iloc[0])]
                                df = df.iloc[1:]
                            else:
                                headers = [f"Column_{i}" for i in range(len(df.columns))]
                        
                        unique_headers = []
                        header_counts = {}
                        
                        for h in headers:
                            if h in header_counts:
                                header_counts[h] += 1
                                unique_headers.append(f"{h}_{header_counts[h]}")
                            else:
                                header_counts[h] = 0
                                unique_headers.append(h)
                        
                        df.columns = unique_headers
                    
                    document_content.append({
                        "page": page_num + 1,
                        "type": "table",
                        "table_number": table_num + 1,
                        "dataframe": df
                    })
    
    return document_content

def format_table_for_llm(df: pd.DataFrame, table_info: dict) -> str:
    if df.empty:
        return f"Empty table on page {table_info['page']}"
    
    table_text = f"\n--- TABLE {table_info['table_number']} (Page {table_info['page']}) ---\n"
    
    table_text += df.to_string(index=False, na_rep='') + "\n"
    
    table_text += "\nKey Financial Data from this table:\n"
    for col in df.columns:
        non_null_values = df[col].dropna()
        if not non_null_values.empty:
            numeric_values = []
            for val in non_null_values:
                if isinstance(val, (int, float)) or (isinstance(val, str) and any(char.isdigit() for char in str(val))):
                    numeric_values.append(str(val))
            
            if numeric_values:
                table_text += f"- {col}: {', '.join(numeric_values[:5])}\n"
    
    table_text += "--- END TABLE ---\n"
    return table_text

def load_customer_pdf_with_vision(file_path):
    document_content = extract_tables_from_pdf(file_path)
    
    page_extraction_quality = {}
    total_meaningful_content = False
    
    for content in document_content:
        page_num = content["page"]
        if content["type"] == "text":
            text_length = len(content["content"].strip())
            page_extraction_quality[page_num] = text_length
            if text_length > 50:
                total_meaningful_content = True
        elif content["type"] == "table" and not content["dataframe"].empty:
            page_extraction_quality[page_num] = page_extraction_quality.get(page_num, 0) + 100
            total_meaningful_content = True
    
    pages_needing_ocr = []
    for page_num, quality in page_extraction_quality.items():
        if quality < 30:
            pages_needing_ocr.append(page_num)
    
    if not total_meaningful_content:
        st.info(f"📸 Document appears to be entirely scanned. Using Google Vision API...")
        vision_documents = extract_text_with_vision(file_path)
        if vision_documents:
            return vision_documents
        else:
            st.warning("Vision API failed, using available content from pdfplumber...")
    
    elif pages_needing_ocr:
        st.info(f"📸 Using OCR for pages with poor text extraction: {pages_needing_ocr}")
        vision_documents = extract_text_with_vision_selective(file_path, pages_needing_ocr)
        
        documents = []
        ocr_pages = {doc.metadata["page"]: doc for doc in vision_documents}
        
        for content in document_content:
            page_num = content["page"]
            
            if page_num in ocr_pages and page_num in pages_needing_ocr:
                documents.append(ocr_pages[page_num])
            else:
                if content["type"] == "text":
                    doc = Document(
                        page_content=content["content"],
                        metadata={
                            "source": file_path,
                            "page": content["page"],
                            "type": "text"
                        }
                    )
                    documents.append(doc)
                elif content["type"] == "table":
                    table_text = format_table_for_llm(content["dataframe"], content)
                    doc = Document(
                        page_content=table_text,
                        metadata={
                            "source": file_path,
                            "page": content["page"],
                            "type": "table",
                            "table_number": content["table_number"]
                        }
                    )
                    documents.append(doc)
        
        return documents
    
    documents = []
    for content in document_content:
        if content["type"] == "text":
            doc = Document(
                page_content=content["content"],
                metadata={
                    "source": file_path,
                    "page": content["page"],
                    "type": "text"
                }
            )
            documents.append(doc)
            
        elif content["type"] == "table":
            table_text = format_table_for_llm(content["dataframe"], content)
            doc = Document(
                page_content=table_text,
                metadata={
                    "source": file_path,
                    "page": content["page"],
                    "type": "table",
                    "table_number": content["table_number"]
                }
            )
            documents.append(doc)
    
    return documents

def extract_text_with_vision_selective(pdf_path, target_pages=None):
    """Extract text using Vision API for specific pages only"""
    vision_client = setup_vision_client()
    if not vision_client:
        return []
    
    try:
        images = pdf_to_images(pdf_path)
        documents = []
        
        for img_info in images:
            if target_pages and img_info["page"] not in target_pages:
                continue
                
            image = vision.Image(content=img_info["data"])
            
            response = vision_client.text_detection(image=image)
            
            if response.error.message:
                st.error(f"Vision API error on page {img_info['page']}: {response.error.message}")
                continue
            
            texts = response.text_annotations
            if texts:
                extracted_text = texts[0].description
                
                doc = Document(
                    page_content=extracted_text,
                    metadata={
                        "source": pdf_path,
                        "page": img_info["page"],
                        "type": "scanned_text",
                        "extraction_method": "google_vision"
                    }
                )
                documents.append(doc)
        
        return documents
        
    except Exception as e:
        st.error(f"Error with Google Vision API: {str(e)}")
        return []

def load_pdf_with_tables(file_path):
    document_content = extract_tables_from_pdf(file_path)
    documents = []
    
    for content in document_content:
        if content["type"] == "text":
            doc = Document(
                page_content=content["content"],
                metadata={
                    "source": file_path,
                    "page": content["page"],
                    "type": "text"
                }
            )
            documents.append(doc)
            
        elif content["type"] == "table":
            table_text = format_table_for_llm(content["dataframe"], content)
            doc = Document(
                page_content=table_text,
                metadata={
                    "source": file_path,
                    "page": content["page"],
                    "type": "table",
                    "table_number": content["table_number"]
                }
            )
            documents.append(doc)
    
    return documents

def split_text(documents):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1500,
        chunk_overlap=200,
        add_start_index=True
    )
    
    chunked_docs = []
    for doc in documents:
        if doc.metadata.get("type") == "table":
            chunked_docs.append(doc)
        else:
            chunks = text_splitter.split_documents([doc])
            chunked_docs.extend(chunks)
    
    return chunked_docs

def extract_financial_info(documents):
    financial_keywords = [
        "salary", "income", "annual income", "monthly income", 
        "basic pay", "gross salary", "CTC",
        "ITR", "income tax return", "form 16", "tax",
        "mutual fund", "SIP", "investment", "portfolio",
        "credit card", "investment amount", "units", "NAV",
        "bank statement", "account balance", "savings",
        "loan", "EMI", "debt", "liability", "credit",
        "bonus", "incentive", "allowance", "deduction",
        "amount", "balance", "value", "total", "sum"
    ]
    
    relevant_chunks = []
    for doc in documents:
        content_lower = doc.page_content.lower()
        
        if doc.metadata.get("type") == "table":
            relevant_chunks.append(doc)
        elif any(keyword in content_lower for keyword in financial_keywords):
            relevant_chunks.append(doc)
    
    return relevant_chunks

comprehensive_template = """
You are an expert Financial Underwriting AI Assistant specialized in insurance policy underwriting. Your primary task is to analyze customer financial documents and calculate financial viability using the EXACT methods specified in the underwriting guidelines for each document type.
Also, you need act as an expert Salary Slip reader, Mutual Fund reader, Bank Statement: Salaried reader, Bank Statement: Closing Balance reader, ITR reader, Form 16 reader and Credit Card reader.

**CRITICAL WORKFLOW - FOLLOW THESE STEPS IN ORDER:**

**DOCUMENT TYPE IDENTIFICATION**
Analyze the customer documents and identify the PRIMARY document type from these categories:
- Salary Slips
- Bank Statement (Salaried)
- Bank Statement (Closing Balance)
- ITR & COI (Income Tax Return & Certificate of Income)
- Form 16
- Mutual Fund Statement - SIP
- Credit Card Statements
- Car Ownership Documents
- Fixed Deposits
- Home Loan
- House/Shop Ownership

**SPECIAL BANK STATEMENT DETECTION RULES:**
For bank statements, you MUST first determine which calculation method to use:

1. **SALARY DETECTION PHASE:** Carefully scan the bank statement for salary-related transactions:
   - Look for regular monthly credits with terms like: "SALARY", "SAL", "PAY", "PAYROLL", "WAGES", "MONTHLY CREDIT"

2. CLOSING BALANCE DETECTION:
   - Carefully scan the bank statement for the balance at the end of each month
   - Extract the balance remaining post all debit-credit for the 3 latest months
   
2. **AUTOMATIC METHOD SELECTION:**
   - IF salary credits are found: Use "Bank Statement (Salaried)" method
   - IF NO salary credits found: Use "Bank Statement (Closing Balance)" method

**Output Format:**
```
PRIMARY DOCUMENT TYPE IDENTIFIED: [Document Type]
BANK STATEMENT SUB-TYPE: [Salaried/Closing Balance] (only for bank statements)
CONFIDENCE LEVEL: High/Medium/Low
SUPPORTING EVIDENCE: [Key identifiers found in the document]
```

**GUIDELINE LOOKUP AND FORMULA EXTRACTION**
Based on the identified document type, customer age ({customer_age}), and policy type ({policy_type}), extract the EXACT calculation method from the guidelines.

**From the Guidelines Document, use these SPECIFIC formulas:**

**For Salary Slips:**
- Term Cases: Annual Salary = Gross Monthly Salary (latest month) × 12; Financial Viability = Annual Salary × Income Multiplier
- Non-Term Cases: Annual Salary = Gross Monthly Salary (latest month) × 12; Annual Bonus = Annual Salary × 0.10; Total = Annual Salary + Annual Bonus; Financial Viability = Total × Income Multiplier

IMPORTANT: For Salary Slips, always consider "Gross Salary".

**For Bank Statement (Salaried) - ONLY when salary credits are detected:**
- Term Cases: Average Monthly Salary (last 3 months) × 12; Add 30% to get Gross Annual Salary; Financial Viability = Gross Annual Salary × Age-based Multiplier
- Non-Term Cases: Average Monthly Salary (last 6 months) × 12; Financial Viability = Annual Income × Age-based Multiplier

**For Bank Statement (Closing Balance) - ONLY when NO salary credits found:**
- Term Cases: Average Closing Balance (last 3 months) × 12; Financial Viability = Annual Average Income × Age-based Multiplier
- Non-Term Cases: Same as Term Cases
IMPORTANT: CLOSING BALANCE IS CONSIDERED TO BE THE AVAILABLE BALANCE FOR A GIVEN MONTH ACCORDING TO THE LATEST DATE. YOU NEED TO CONSIDER THE LAST AVAILABLE BALANCE FOR THE 3 LATEST MONTHS IN THE DOCUMENTS.

**For ITR & COI:**
- Term Cases: Only Earned Income (exclude unearned); Financial Viability = Total Earned Income × Age-based Multiplier
- Non-Term Cases: Include ALL income types (earned + unearned); Financial Viability = Total Income × Age-based Multiplier

**For Form 16:**
- Term Cases: Gross Income from Part A; Financial Viability = Annual Income × Age-based Multiplier

**For Mutual Fund Statement - SIP:**
- Term Cases: Monthly SIP × 12 = Annual Income; Financial Viability = Annual Income × Age-based Multiplier

**For Credit Card Statements:**
- Term Cases: Monthly CC Statement Value × 6 = Annual Income; Financial Viability = Annual Income × Age-based Multiplier

**For Car Ownership:**
- Term Cases: Car IDV Value × 2 = Annual Income; Financial Viability = Annual Income × Age-based Multiplier

**For Fixed Deposits:**
- Term Cases: Investment Value × 0.05 = Estimated Annual Income; Financial Viability = Annual Income × Age-based Multiplier

**For Home Loan:**
- Term Cases: Monthly EMI × 24 = Annual Income; Financial Viability = Annual Income × Age-based Multiplier

**For House/Shop Ownership:**
- Term Cases: Financial Viability = Property Value × 0.50

**AGE-BASED MULTIPLIER SELECTION**
Use the correct age-based multiplier from guidelines:

**Term Cases Multipliers:**
- Age 18-30: 25x
- Age 31-35: 25x  
- Age 36-40: 20x
- Age 41-45: 15x
- Age 46-50: 12x
- Age 51-55: 10x
- Age ≥56: 5x

**Non-Term Cases Multipliers:**
- Age 18-30: 35x
- Age 31-35: 30x
- Age 36-40: 25x
- Age 41-45: 20x
- Age 46-50: 15x
- Age 51-65: 10x
- Age >65: 6x

**PRECISE CALCULATION**
Apply the exact formula for the identified document type using actual values from customer documents.

**For Bank Statements - Show Detection Process:**
```
BANK STATEMENT ANALYSIS:
Salary Detection Scan:
- Searched for: [List salary-related keywords found/not found]
- Regular Monthly Credits: [Found/Not Found]
- Salary Patterns: [Describe any patterns found]
- Decision: [Salaried Method / Closing Balance Method]

CALCULATION METHOD SELECTED: [Method name]
```

**Output Format:**
```
GUIDELINE-BASED CALCULATION:
Document Type: [Type]
Sub-type: [If bank statement, specify Salaried or Closing Balance]
Policy Type: {policy_type}
Customer Age: {customer_age}
Applicable Formula: [Exact formula from guidelines]
Age-based Multiplier: [X]x

INPUT VALUES:
[List all extracted values with source references]

CALCULATION STEPS:
Step 1: [First calculation with actual numbers]
Step 2: [Second calculation if applicable]
Step 3: [Final calculation]

FINANCIAL VIABILITY: ₹[Final Amount]
METHOD JUSTIFICATION: [Why this specific method was chosen for bank statements]
```

**VALIDATION AND COMPARISON**
Compare your calculation with the generic method to ensure accuracy.

**SUMMARY REPORT**
Provide detailed analysis including:

**Document Analysis Summary:**
| Parameter | Value | Document Type Method |
|-----------|-------|---------------------|

This table should hold Document type, Customer Age, Policy Type and the important financials extracted from the uploaded documents.

**Risk Assessment:**
- Income Stability: [Analysis]
- Debt-to-Income Ratio: [If applicable]
- Financial Capacity: [Assessment]
- Premium Affordability: [Based on calculated viability]

**Final Recommendation:**
- Policy Eligibility: Approved/Conditional/Declined
- Justification: [Why this specific calculation method was used]
- What all other documents are required to judge the financial viability better

Also, in one or two lines, give a brief conclusion on the customer's overall financial status

**CRITICAL INSTRUCTIONS:**
1. NEVER use the generic age-based multiplier ({income_multiplier}x) alone - it's only for reference
2. ALWAYS use the document-type-specific calculation method from guidelines
3. For bank statements, ALWAYS perform salary detection first, then choose appropriate method
4. Extract EXACT numerical values from customer documents
5. Show ALL calculation steps with actual numbers
6. Reference specific guideline sections/pages
7. If document type is unclear, state this and explain your reasoning
8. If multiple document types are present, prioritize the most reliable one (usually Salary Slip or ITR)

**Customer Information:**
- Age: {customer_age} years
- Policy Type: {policy_type}
- Reference Generic Multiplier: {income_multiplier}x (DO NOT USE - for reference only)

**SUMMARY REPORT**
- Give a detailed summarized report for the customer and the uploaded document with a good recommendation his/her financial viability

**Question:** {question}
**Guidelines Context:** {guidelines_context}
**Customer Financial Documents:** {customer_context}

**RESPONSE:**
"""

specific_template = """
You are a financial underwriting expert. Answer the specific question asked based on the customer's financial documents and underwriting guidelines. 
IMPORTANT: Mention the customer financial document type.
IMPORTANT: PLEASE DO CORRECT CALCULATIONS
CRITICAL: All NUMBERS TO BE CONSIDERED IN INDIAN RUPEES AND SYSTEM

**BANK STATEMENT SPECIAL HANDLING:**
If analyzing bank statements, you MUST:

1. **FIRST: Detect Salary Credits**
   - Scan for salary-related transactions: "SALARY", "SAL", "PAY", "PAYROLL", "WAGES"
   - Look for regular monthly employer credits
   - Check transaction patterns over 3-6 months

2. **THEN: Choose Calculation Method**
   - IF salary credits found → Use Bank Statement (Salaried) method
   - IF NO salary credits found → Use Bank Statement (Closing Balance) method

3. **SHOW YOUR DETECTION PROCESS**
   ```
   BANK STATEMENT DETECTION:
   Salary Search Result: [Found/Not Found]
   Evidence: [List specific transactions or lack thereof]
   Method Selected: [Salaried/Closing Balance]
   ```

**Customer Information:**
- Age: {customer_age} years
- Policy Type: {policy_type}
- Income Multiplier: {income_multiplier}x

Based on the identified document type, customer age ({customer_age}), and policy type ({policy_type}), extract the EXACT calculation method from the guidelines.

**From the Guidelines Document, use these SPECIFIC formulas:**

**For Salary Slips:**
- Term Cases: Annual Salary = Monthly Salary × 12; Financial Viability = Annual Salary × Income Multiplier
- Non-Term Cases: Annual Salary = Monthly Salary × 12; Annual Bonus = Annual Salary × 0.10; Total = Annual Salary + Annual Bonus; Financial Viability = Total × Income Multiplier

**For Bank Statement (Salaried) - USE ONLY when salary credits detected:**
- Term Cases: Average Monthly Salary (last 3 months) × 12; Add 30% to get Gross Annual Salary; Financial Viability = Gross Annual Salary × Age-based Multiplier
- Non-Term Cases: Average Monthly Salary (last 6 months) × 12; Financial Viability = Annual Income × Age-based Multiplier

**For Bank Statement (Closing Balance) - USE ONLY when NO salary credits found:**
- Term Cases: Average Closing Balance (last 3 months) × 12; Financial Viability = Annual Average Income × Age-based Multiplier
- Non-Term Cases: Same as Term Cases

**For ITR & COI:**
- Term Cases: Only Earned Income (exclude unearned); Financial Viability = Total Earned Income × Age-based Multiplier
- Non-Term Cases: Include ALL income types (earned + unearned); Financial Viability = Total Income × Age-based Multiplier

**For Form 16:**
- Term Cases: Gross Income from Part A; Financial Viability = Annual Income × Age-based Multiplier

**For Mutual Fund Statement - SIP:**
- Term Cases: Monthly SIP × 12 = Annual Income; Financial Viability = Annual Income × Age-based Multiplier

**For Credit Card Statements:**
- Term Cases: Monthly CC Statement Value × 6 = Annual Income; Financial Viability = Annual Income × Age-based Multiplier

**For Car Ownership:**
- Term Cases: Car IDV Value × 2 = Annual Income; Financial Viability = Annual Income × Age-based Multiplier

**For Fixed Deposits:**
- Term Cases: Investment Value × 0.05 = Estimated Annual Income; Financial Viability = Annual Income × Age-based Multiplier

**For Home Loan:**
- Term Cases: Monthly EMI × 24 = Annual Income; Financial Viability = Annual Income × Age-based Multiplier

**For House/Shop Ownership:**
- Term Cases: Financial Viability = Property Value × 0.50

**AGE-BASED MULTIPLIER SELECTION**
Use the correct age-based multiplier from guidelines:

**Term Cases Multipliers:**
- Age 18-30: 25x
- Age 31-35: 25x  
- Age 36-40: 20x
- Age 41-45: 15x
- Age 46-50: 12x
- Age 51-55: 10x
- Age ≥56: 5x

**Non-Term Cases Multipliers:**
- Age 18-30: 35x
- Age 31-35: 30x
- Age 36-40: 25x
- Age 41-45: 20x
- Age 46-50: 15x
- Age 51-65: 10x
- Age >65: 6x

**PRECISE CALCULATION**
Apply the exact formula for the identified document type using actual values from customer documents.

**Output Format:**
```
GUIDELINE-BASED CALCULATION:
Document Type: [Type]
Policy Type: {policy_type}
Customer Age: {customer_age}
Applicable Formula: [Exact formula from guidelines]
Age-based Multiplier: [X]x

INPUT VALUES:
[List all extracted values with source references]

CALCULATION STEPS:
Step 1: [First calculation with actual numbers]
Step 2: [Second calculation if applicable]
Step 3: [Final calculation]

FINANCIAL VIABILITY: ₹[Final Amount]
```

**VALIDATION AND COMPARISON**
Compare your calculation with the generic method to ensure accuracy.

**COMPREHENSIVE ANALYSIS**
Provide detailed analysis including:

**Document Analysis Summary:**
| Parameter | Value | Source Location | Document Type Method | Age-Generic Method | Variance |
|-----------|-------|----------------|---------------------|-------------------|----------|

**Risk Assessment:**
- Income Stability: [Analysis]
- Debt-to-Income Ratio: [If applicable]
- Financial Capacity: [Assessment]
- Premium Affordability: [Based on calculated viability]

**Final Recommendation:**
- Policy Eligibility: Approved/Conditional/Declined
- Justification: [Why this specific calculation method was used]

**SUMMARY REPORT**
- Give a detailed summarized report for the customer and the uploaded document with a good recommendation his/her financial viability

**CRITICAL INSTRUCTIONS:**
1. NEVER use the generic age-based multiplier ({income_multiplier}x) alone - it's only for reference
2. ALWAYS use the document-type-specific calculation method from guidelines
3. Extract EXACT numerical values from customer documents
4. Show ALL calculation steps with actual numbers
5. Reference specific guideline sections/pages
6. If document type is unclear, state this and explain your reasoning
7. If multiple document types are present, prioritize the most reliable one (usually Salary Slip or ITR)

**Customer Information:**
- Age: {customer_age} years
- Policy Type: {policy_type}
- Reference Generic Multiplier: {income_multiplier}x (DO NOT USE - for reference only)

**Question:** {question}
**Guidelines Context:** {guidelines_context}
**Customer Financial Documents:** {customer_context}

**RESPONSE:**
"""

specific_template = """
You are a financial underwriting expert. Answer the specific question asked based on the customer's financial documents and underwriting guidelines. 
IMPORTANT: Mention the customer financial document type.
IMPORTANT: PLEASE DO CORRECT CALCULATIONS
CRITICAL: All NUMBERS TO BE CONSIDERED IN INDIAN RUPEES AND SYSTEM


**Customer Information:**
- Age: {customer_age} years
- Policy Type: {policy_type}
- Income Multiplier: {income_multiplier}x

Based on the identified document type, customer age ({customer_age}), and policy type ({policy_type}), extract the EXACT calculation method from the guidelines.

**From the Guidelines Document, use these SPECIFIC formulas:**

**For Salary Slips:**
- Term Cases: Annual Salary = Gross Monthly Salary × 12; Financial Viability = Annual Salary × Income Multiplier
- Non-Term Cases: Annual Salary = Gross Monthly Salary × 12; Annual Bonus = Annual Salary × 0.10; Total = Annual Salary + Annual Bonus; Financial Viability = Total × Income Multiplier

**For Bank Statement (Salaried):**
- Term Cases: Average Monthly Salary (last 3 months) × 12; Add 30% to get Gross Annual Salary; Financial Viability = Gross Annual Salary × Age-based Multiplier
- Non-Term Cases: Average Monthly Salary (last 6 months) × 12; Financial Viability = Annual Income × Age-based Multiplier

**For Bank Statement (Closing Balance):**
- Term Cases: Average Closing Balance (last 3 months) × 12; Financial Viability = Annual Average Income × Age-based Multiplier

**For ITR & COI:**
- Term Cases: Only Earned Income (exclude unearned); Financial Viability = Total Earned Income × Age-based Multiplier
- Non-Term Cases: Include ALL income types (earned + unearned); Financial Viability = Total Income × Age-based Multiplier

**For Form 16:**
- Term Cases: Gross Income from Part A; Financial Viability = Annual Income × Age-based Multiplier

**For Mutual Fund Statement - SIP:**
- Term Cases: Monthly SIP × 12 = Annual Income; Financial Viability = Annual Income × Age-based Multiplier

**For Credit Card Statements:**
- Term Cases: Monthly CC Statement Value × 6 = Annual Income; Financial Viability = Annual Income × Age-based Multiplier

**For Car Ownership:**
- Term Cases: Car IDV Value × 2 = Annual Income; Financial Viability = Annual Income × Age-based Multiplier

**For Fixed Deposits:**
- Term Cases: Investment Value × 0.05 = Estimated Annual Income; Financial Viability = Annual Income × Age-based Multiplier

**For Home Loan:**
- Term Cases: Monthly EMI × 24 = Annual Income; Financial Viability = Annual Income × Age-based Multiplier

**For House/Shop Ownership:**
- Term Cases: Financial Viability = Property Value × 0.50

AGE-BASED MULTIPLIER SELECTION**
Use the correct age-based multiplier from guidelines:

**Term Cases Multipliers:**
- Age 18-30: 25x
- Age 31-35: 25x  
- Age 36-40: 20x
- Age 41-45: 15x
- Age 46-50: 12x
- Age 51-55: 10x
- Age ≥56: 5x

**Non-Term Cases Multipliers:**
- Age 18-30: 35x
- Age 31-35: 30x
- Age 36-40: 25x
- Age 41-45: 20x
- Age 46-50: 15x
- Age 51-65: 10x
- Age >65: 6x

PRECISE CALCULATION**
Apply the exact formula for the identified document type using actual values from customer documents.

IMPORTANT: The documents contain both TEXT and TABLE data.

CRITICAL: Please search carefully in both text content and tabular data. Financial documents often have key information in table format. Also, do not mention steps in the output.

Be concise and specific. Only provide the information directly relevant to the question asked. If you find the information in a table, mention the table number and page.

Question: {question}
Customer Financial Documents: {customer_context}
Answer:
"""

def determine_question_type(question: str) -> str:
    comprehensive_keywords = [
        "complete analysis", "full report", "comprehensive", "detailed analysis",
        "overall assessment", "complete evaluation", "full evaluation",
        "risk assessment", "financial capacity", "policy eligibility"
    ]
    
    question_lower = question.lower()
    
    if any(keyword in question_lower for keyword in comprehensive_keywords):
        return "comprehensive"
    
    specific_indicators = [
        "what is", "how much", "when", "where", "which", "who",
        "calculate", "show me", "find", "extract", "tell me about"
    ]
    
    if any(indicator in question_lower for indicator in specific_indicators):
        return "specific"
    
    return "specific"

guidelines_directory = '.github/guidelines/'
customer_docs_directory = '.github/customer_docs/'

os.makedirs(guidelines_directory, exist_ok=True)
os.makedirs(customer_docs_directory, exist_ok=True)

embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
guidelines_vector_store = InMemoryVectorStore(embeddings)
customer_docs_vector_store = InMemoryVectorStore(embeddings)

def initialize_model():
    """Initialize the ChatGroq model with API key from session state"""
    if not st.session_state.groq_api_key:
        return None
    
    try:
        model = ChatGroq(
            groq_api_key=st.session_state.groq_api_key,
            model_name="meta-llama/llama-4-scout-17b-16e-instruct", 
            temperature=0.3
        )
        st.session_state.model_initialized = True
        return model
    except Exception as e:
        st.error(f"Error initializing Groq model: {str(e)}")
        return None

def upload_pdf(file, directory):
    file_path = directory + file.name
    with open(file_path, "wb") as f:
        f.write(file.getbuffer())
    return file_path

def process_documents_with_pii_shield(documents):
    protected_docs = []
    for doc in documents:
        anonymized_content = pii_shield.anonymize_text(doc.page_content)
        
        protected_doc = Document(
            page_content=anonymized_content,
            metadata=doc.metadata
        )
        protected_docs.append(protected_doc)
    
    return protected_docs

def analyze_customer_finances(question, guidelines_docs, customer_docs):
    if not st.session_state.groq_api_key:
        st.error("⚠️ Please enter your Groq API key to proceed with analysis.")
        return "API key required for analysis."
    
    if not st.session_state.model_initialized:
        model = initialize_model()
        if not model:
            return "Failed to initialize model. Please check your API key."
    else:
        model = initialize_model()
    guidelines_context = "\n\n".join([doc.page_content for doc in guidelines_docs])
    
    financial_docs = extract_financial_info(customer_docs)
    customer_context = "\n\n".join([doc.page_content for doc in financial_docs])
    
    question_type = determine_question_type(question)
    
    customer_age = st.session_state.customer_age or 30
    policy_type = st.session_state.policy_type or "Term"
    income_multiplier = get_income_multiplier(customer_age, policy_type)
    
    if question_type == "comprehensive":
        template = comprehensive_template
    else:
        template = specific_template
    
    prompt = ChatPromptTemplate.from_template(template)
    chain = prompt | model
    
    response = chain.invoke({
        "question": question,
        "guidelines_context": guidelines_context,
        "customer_context": customer_context,
        "customer_age": customer_age,
        "policy_type": policy_type,
        "income_multiplier": income_multiplier
    })
    
    return response.content

if "conversation_history" not in st.session_state:
    st.session_state.conversation_history = []
if "guidelines_loaded" not in st.session_state:
    st.session_state.guidelines_loaded = False
if "customer_docs_loaded" not in st.session_state:
    st.session_state.customer_docs_loaded = False
if "table_stats" not in st.session_state:
    st.session_state.table_stats = {"total_tables": 0, "tables_by_page": {}}
if "extracted_financial_data" not in st.session_state:
    st.session_state.extracted_financial_data = None
if "customer_age" not in st.session_state:
    st.session_state.customer_age = None
if "policy_type" not in st.session_state:
    st.session_state.policy_type = "Term"

if "google_vision_api_key" not in st.session_state:
    st.session_state.google_vision_api_key = ""
if "groq_api_key" not in st.session_state:
    st.session_state.groq_api_key = ""
if "model_initialized" not in st.session_state:
    st.session_state.model_initialized = False

if "last_risk_assessment" not in st.session_state:
    st.session_state.last_risk_assessment = False
    

with st.sidebar:
    pii_enabled = st.toggle("Enable PII Shield", value=True, help="Automatically anonymize personal information")
    pii_shield.anonymization_enabled = pii_enabled
    
    if pii_enabled:
        st.success("🛡️ PII Shield Active")
        
        if pii_shield.replacement_map:
            st.markdown("#### PII Detection Summary")
            pii_summary = pii_shield.get_pii_summary()
            for pii_type, count in pii_summary.items():
                st.write(f"• {pii_type.replace('_', ' ').title()}: {count} instances")
    else:
        st.warning("⚠️ PII Shield Disabled - Use with caution!")
    st.markdown("---")
    st.markdown("### 🔑 API Configuration")
    
    groq_key = st.text_input(
        "Groq API Key",
        type="password",
        value=st.session_state.groq_api_key,
        help="Enter your Groq API key for LLM processing"
    )
    
    if groq_key != st.session_state.groq_api_key:
        st.session_state.groq_api_key = groq_key
        st.session_state.model_initialized = False
    
    vision_key = st.text_input(
        "Google Vision API Key",
        type="password",
        value=st.session_state.google_vision_api_key,
        help="Enter your Google Vision API key for OCR processing"
    )
    
    if vision_key != st.session_state.google_vision_api_key:
        st.session_state.google_vision_api_key = vision_key
    
    if st.session_state.groq_api_key:
        st.success("✅ Groq API Key Set")
    else:
        st.error("❌ Groq API Key Required")
    
    if st.session_state.google_vision_api_key:
        st.success("✅ Google Vision API Key Set")
    else:
        st.warning("⚠️ Google Vision API Key Optional (for scanned PDFs)")
    st.markdown("### 📸 Google Vision Status")
    vision_client = setup_vision_client()
    if vision_client:
        st.success("✅ Google Vision API Ready")
    else:
        st.error("❌ Google Vision API Not Available")
        st.markdown("Set `GOOGLE_VISION_API_KEY` in secrets.toml or environment")
    
    if st.button("🗑️ Clear Analysis History"):
        st.session_state.conversation_history = []
        st.session_state.table_stats = {"total_tables": 0, "tables_by_page": {}}
        pii_shield.replacement_map.clear()
        st.rerun()
    
    if st.button("🧹 Clear PII Cache"):
        pii_shield.replacement_map.clear()
        st.success("PII cache cleared")

col1, col2 = st.columns(2)

with col1:
    guidelines_files = st.file_uploader(
        "📋 Upload Financial Underwriting Guidelines",
        type="pdf",
        accept_multiple_files=True,
        key="guidelines_uploader"
    )
    
    if guidelines_files and not st.session_state.guidelines_loaded:
        with st.spinner("Processing guidelines with table extraction..."):
            all_guideline_docs = []
            for file in guidelines_files:
                file_path = upload_pdf(file, guidelines_directory)
                documents = load_pdf_with_tables(file_path)
                chunked_documents = split_text(documents)
                all_guideline_docs.extend(chunked_documents)
            
            guidelines_vector_store.add_documents(all_guideline_docs)
            st.session_state.guidelines_loaded = True
            st.success(f"✅ {len(guidelines_files)} guideline document(s) processed successfully!")

with col2:    
    customer_files = st.file_uploader(
        "💼 Upload Customer Financial Documents",
        type="pdf",
        accept_multiple_files=True,
        key="customer_uploader",
        help="Supports both regular PDFs and scanned documents (using Google Vision API)"
    )
    
    if customer_files:
        with st.spinner("Processing customer documents with enhanced table extraction and Vision API..."):
            all_customer_docs = []
            table_count = 0
            tables_by_page = {}
            scanned_count = 0
            
            customer_doc_names = [file.name for file in customer_files]
            st.session_state.customer_doc_names = customer_doc_names
            
            for file in customer_files:
                file_path = upload_pdf(file, customer_docs_directory)
                
                documents = load_customer_pdf_with_vision(file_path)
                
                if any(doc.metadata.get("extraction_method") == "google_vision" for doc in documents):
                    scanned_count += 1
                
                for doc in documents:
                    if doc.metadata.get("type") == "table":
                        table_count += 1
                        page = doc.metadata.get("page", 0)
                        tables_by_page[page] = tables_by_page.get(page, 0) + 1
                
                chunked_documents = split_text(documents)
                
                if pii_shield.anonymization_enabled:
                    protected_documents = process_documents_with_pii_shield(chunked_documents)
                    all_customer_docs.extend(protected_documents)
                else:
                    all_customer_docs.extend(chunked_documents)
            
            st.session_state.table_stats = {
                "total_tables": table_count,
                "tables_by_page": tables_by_page
            }
            
            customer_docs_vector_store.add_documents(all_customer_docs)
            st.session_state.customer_docs_loaded = True
            
            success_msg = f"✅ {len(customer_files)} customer document(s) processed!"
            if scanned_count > 0:
                success_msg += f" 📸 {scanned_count} scanned with Vision API!"
            if table_count > 0:
                success_msg += f" 📊 {table_count} tables extracted!"
            
            if pii_shield.anonymization_enabled and pii_shield.replacement_map:
                success_msg += f" 🛡️ {len(pii_shield.replacement_map)} PII elements anonymized!"
            
            st.success(success_msg)

if st.session_state.guidelines_loaded and st.session_state.customer_docs_loaded:
    st.success("🎉 All documents loaded! Enhanced table extraction and Vision API ready for financial analysis.")
elif st.session_state.guidelines_loaded:
    st.warning("Guidelines loaded. Please upload customer financial documents.")
elif st.session_state.customer_docs_loaded:
    st.warning("Customer documents loaded. Please upload underwriting guidelines.")
else:
    st.info("📤 Please upload both guidelines and customer financial documents to begin analysis.")

if st.session_state.guidelines_loaded and st.session_state.customer_docs_loaded:
    st.markdown("---")
    st.markdown("### 👤 Customer Information")
    
    col1, col2 = st.columns(2)
    
    with col1:
        customer_age = st.number_input(
            "Customer Age",
            min_value=18,
            max_value=80,
            value=st.session_state.customer_age or 30,
            help="Required for calculating income multiplier"
        )
        st.session_state.customer_age = customer_age
    
    with col2:
        policy_type = st.selectbox(
            "Policy Type",
            ["Term", "Non-Term"],
            index=0 if st.session_state.policy_type == "Term" else 1,
            help="Policy type affects income multiplier calculation"
        )
        st.session_state.policy_type = policy_type
    
    if customer_age:
        multiplier = get_income_multiplier(customer_age, policy_type)
        st.info(f"📊 Current Income Multiplier: **{multiplier}x** (Age: {customer_age}, Policy: {policy_type})")

def create_risk_assessment_docx(assessment_content, customer_age, policy_type, filename="risk_assessment_report.docx"):
    """Create a professionally formatted DOCX document for Risk Assessment with proper markdown parsing"""
 
    doc = DocxDocument()
    
    title = doc.add_heading('Financial Risk Assessment Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('Customer Information', level=1)
    customer_info = doc.add_paragraph()
    customer_info.add_run(f'Customer Age: ').bold = True
    customer_info.add_run(f'{customer_age} years\n')
    customer_info.add_run(f'Policy Type: ').bold = True
    customer_info.add_run(f'{policy_type}\n')    
    doc.add_heading('Risk Assessment Analysis', level=1)
    
    lines = assessment_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
        
        if '|' in line and detect_table_start(line):
            table_lines = []
            j = i
            
            while j < len(lines) and ('|' in lines[j] or lines[j].strip() == '' or lines[j].strip().startswith('---')):
                if lines[j].strip() and not lines[j].strip().startswith('---'):
                    table_lines.append(lines[j].strip())
                j += 1
            
            if table_lines:
                create_docx_table(doc, table_lines)
            
            i = j
            continue
        
        if line.startswith('###'):
            clean_heading = clean_markdown_text(line[3:].strip())
            doc.add_heading(clean_heading, level=3)
        elif line.startswith('##'):
            clean_heading = clean_markdown_text(line[2:].strip())
            doc.add_heading(clean_heading, level=2)
        elif line.startswith('#'):
            clean_heading = clean_markdown_text(line[1:].strip())
            doc.add_heading(clean_heading, level=1)
        elif is_heading_by_content(line):
            clean_heading = clean_markdown_text(line)
            doc.add_heading(clean_heading, level=2)
        elif line.startswith('•') or line.startswith('-') or line.startswith('*'):
            clean_text = clean_markdown_text(line[1:].strip())
            add_formatted_paragraph(doc, clean_text, style='List Bullet')
        elif ':' in line and len(line.split(':', 1)) == 2 and not line.count(':') > 3:
            key, value = line.split(':', 1)
            if len(key.strip()) < 50:
                p = doc.add_paragraph()
                clean_key = clean_markdown_text(key.strip())
                clean_value = clean_markdown_text(value.strip())
                p.add_run(f'{clean_key}: ').bold = True
                add_formatted_text_to_paragraph(p, clean_value)
            else:
                add_formatted_paragraph(doc, clean_markdown_text(line))
        else:
            add_formatted_paragraph(doc, clean_markdown_text(line))
        
        i += 1
    
    doc.add_page_break()
    footer_section = doc.sections[0]
    footer = footer_section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Financial Underwriting Assistant - Confidential Report"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

def clean_markdown_text(text):
    """Remove markdown formatting symbols from text"""
    if not text:
        return ""
    
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'`(.*?)`', r'\1', text)
    text = re.sub(r'#{1,6}\s*', '', text)
    
    return text.strip()

def add_formatted_paragraph(doc, text, style=None):
    """Add a paragraph with markdown formatting converted to DOCX formatting"""
    if not text.strip():
        return
    
    p = doc.add_paragraph(style=style)
    add_formatted_text_to_paragraph(p, text)

def add_formatted_text_to_paragraph(paragraph, text):
    """Add formatted text to a paragraph, converting markdown to DOCX formatting"""
    if not text:
        return
    
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            bold_text = part[2:-2]
            run = paragraph.add_run(bold_text)
            run.bold = True
        else:
            italic_parts = re.split(r'(\*.*?\*)', part)
            for italic_part in italic_parts:
                if italic_part.startswith('*') and italic_part.endswith('*') and not italic_part.startswith('**'):
                    italic_text = italic_part[1:-1]
                    run = paragraph.add_run(italic_text)
                    run.italic = True
                else:
                    if italic_part:
                        paragraph.add_run(italic_part)

def detect_table_start(line):
    """Detect if a line is the start of a markdown-style table"""
    pipe_count = line.count('|')
    if pipe_count >= 2:
        cells = [cell.strip() for cell in line.strip('|').split('|')]
        return len(cells) >= 2 and any(len(cell.strip()) > 0 for cell in cells)
    return False

def is_heading_by_content(line):
    """Determine if a line should be treated as a heading based on content"""
    heading_keywords = [
        'ANALYSIS', 'ASSESSMENT', 'SUMMARY', 'RECOMMENDATION', 'VIABILITY', 
        'CALCULATION', 'DOCUMENT', 'GUIDELINE', 'INPUT VALUES', 'STEPS',
        'RISK', 'FINAL', 'CONCLUSION', 'FINDINGS', 'OVERVIEW'
    ]
    
    line_upper = line.upper()
    
    if line.startswith('#') or '**' in line:
        return False
    
    if line.isupper() and len(line) > 3 and len(line) < 100:
        return True
    
    if any(keyword in line_upper for keyword in heading_keywords):
        return True
    
    if line.endswith(':') and len(line) < 100 and not ':' in line[:-1]:
        return True
    
    return False

def create_docx_table(doc, table_lines):
    """Create a properly formatted DOCX table from markdown-style table lines"""
    if not table_lines:
        return
    
    parsed_rows = []
    for line in table_lines:
        cells = [clean_markdown_text(cell.strip()) for cell in line.strip('|').split('|')]
        if cells and any(cell.strip() for cell in cells):
            parsed_rows.append(cells)
    
    if not parsed_rows:
        return
    
    max_cols = max(len(row) for row in parsed_rows)
    
    for row in parsed_rows:
        while len(row) < max_cols:
            row.append('')
    
    table = doc.add_table(rows=len(parsed_rows), cols=max_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for row_idx, row_data in enumerate(parsed_rows):
        table_row = table.rows[row_idx]
        for col_idx, cell_data in enumerate(row_data):
            cell = table_row.cells[col_idx]
            cell.text = cell_data.strip()
            
            if row_idx == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                
                try:
                    cell_elem = cell._element
                    cell_properties = cell_elem.get_or_add_tcPr()
                    shade_elem = OxmlElement('w:shd')
                    shade_elem.set(qn('w:fill'), 'D9E2F3')
                    cell_properties.append(shade_elem)
                except:
                    pass
    
    for col in table.columns:
        col.width = Inches(2.0)
    
    doc.add_paragraph('')
    
def enhanced_risk_assessment_button():
    """Enhanced Risk Assessment with improved DOCX export option"""
    col_assess, col_export = st.columns([3, 1])
    
    with col_assess:
        risk_assess_clicked = st.button("⚖️ Risk Assessment", use_container_width=True)
    
    with col_export:
        export_clicked = st.button("📄 Export DOCX", use_container_width=True)
    
    if risk_assess_clicked:
        st.session_state.conversation_history.append({
            "role": "user", 
            "content": "Provide a comprehensive financial risk assessment using all available text and tabular data."
        })
        st.session_state.last_risk_assessment = True
    
    if export_clicked:
        risk_assessment_content = None
        for msg in reversed(st.session_state.conversation_history):
            if (msg.get("role") == "assistant" and 
                ("risk assessment" in msg.get("content", "").lower() or 
                 "financial viability" in msg.get("content", "").lower() or
                 "analysis" in msg.get("content", "").lower())):
                risk_assessment_content = msg.get("content")
                break
        
        if risk_assessment_content:
            try:
                docx_buffer = create_risk_assessment_docx(
                    risk_assessment_content, 
                    st.session_state.customer_age or 30, 
                    st.session_state.policy_type or "Term"
                )
                
                customer_doc_name = ""
                if hasattr(st.session_state, 'customer_doc_names') and st.session_state.customer_doc_names:
                    first_doc = st.session_state.customer_doc_names[0]
                    customer_doc_name = first_doc.replace('.pdf', '').replace(' ', '_')
                    if len(customer_doc_name) > 20:
                        customer_doc_name = customer_doc_name[:20]
                    customer_doc_name = f"_{customer_doc_name}"
                
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                filename = f"risk_assessment_report{customer_doc_name}.docx"
                
                st.download_button(
                    label="📥 Download Risk Assessment Report",
                    data=docx_buffer.getvalue(),
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
                st.success("✅ DOCX report generated successfully with proper table formatting!")
                
            except Exception as e:
                st.error(f"Error generating DOCX: {str(e)}")
                st.error("Please check the console for detailed error information.")
                st.write("Debug info:")
                st.write(f"Content length: {len(risk_assessment_content) if risk_assessment_content else 0}")
                st.write(f"Contains tables: {'|' in risk_assessment_content if risk_assessment_content else False}")
        else:
            st.warning("⚠️ No risk assessment found to export. Please run a risk assessment first.")

if st.session_state.guidelines_loaded and st.session_state.customer_docs_loaded:
    st.markdown("---")
    st.markdown("### 🔍 Enhanced Financial Analysis with Table Data & Vision OCR")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        if st.button("💰 Income Analysis", use_container_width=True):
            st.session_state.conversation_history.append({
                "role": "user", 
                "content": "What is the customer's monthly income and income sources? Look for both text and tabular data."
            })
    
    with col2:
        if st.button("📊 Investment Analysis", use_container_width=True):
            st.session_state.conversation_history.append({
                "role": "user", 
                "content": "Analyze the customer's investment portfolio from mutual fund statements and investment tables."
            })
    
    with col3:
        enhanced_risk_assessment_button()
    
    with col4:
        if st.button("📋 Full Report", use_container_width=True):
            st.session_state.conversation_history.append({
                "role": "user", 
                "content": "Provide a complete comprehensive financial analysis report using all text and table data."
            })
    
    question = st.chat_input("Please ask a question")
    
    if question:
        st.session_state.conversation_history.append({"role": "user", "content": question})
    
    if st.session_state.conversation_history and st.session_state.conversation_history[-1]["role"] == "user":
        if st.session_state.customer_age is None:
            st.warning("⚠️ Please enter customer age before proceeding with analysis.")
            st.stop()
    
        with st.spinner("Analyzing financial documents with table data..."):
            latest_question = st.session_state.conversation_history[-1]["content"]
            guidelines_docs = guidelines_vector_store.similarity_search(latest_question, k=5)
            customer_docs = customer_docs_vector_store.similarity_search(latest_question, k=15)
        
            answer = analyze_customer_finances(latest_question, guidelines_docs, customer_docs)
        
            st.session_state.conversation_history.append({"role": "assistant", "content": answer})
    
    for message in st.session_state.conversation_history:
        if message["role"] == "user":
            st.chat_message("user").write(message["content"])
        elif message["role"] == "assistant":
            st.chat_message("assistant").write(message["content"])
